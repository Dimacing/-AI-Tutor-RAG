from __future__ import annotations

import hashlib
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from urllib.parse import urldefrag, urljoin, urlparse

import requests
from bs4 import BeautifulSoup

from .chunking import chunk_text, normalize_whitespace


@dataclass(frozen=True)
class SourceItem:
    title: str
    url: str | None
    path: str | None
    crawl: bool


@dataclass(frozen=True)
class Document:
    source_id: str
    title: str
    source_url: str
    text: str


@dataclass(frozen=True)
class Chunk:
    chunk_id: str
    source_id: str
    title: str
    source_url: str
    text: str


def load_sources(sources_file: Path) -> list[SourceItem]:
    raw = json.loads(sources_file.read_text(encoding="utf-8"))
    base_dir = sources_file.parent
    items: list[SourceItem] = []
    for entry in raw:
        path = entry.get("path")
        if path:
            candidate = Path(path)
            if not candidate.is_absolute():
                path = str((base_dir / candidate).resolve())
        items.append(
            SourceItem(
                title=str(entry.get("title", "Untitled")),
                url=entry.get("url"),
                path=path,
                crawl=bool(entry.get("crawl", False)),
            )
        )
    return items


def expand_sources(
    sources: Iterable[SourceItem],
    allowlist: str,
    max_pages: int,
) -> list[SourceItem]:
    expanded: list[SourceItem] = []
    visited: set[str] = set()

    for item in sources:
        expanded.append(item)
        if item.url:
            visited.add(item.url)

        if item.crawl and item.url:
            try:
                links = discover_links(item.url, allowlist, max_pages)
            except Exception:
                links = []
            for link in links:
                if link in visited:
                    continue
                visited.add(link)
                expanded.append(
                    SourceItem(title=link, url=link, path=None, crawl=False)
                )

    return expanded


def discover_links(url: str, allowlist: str, max_pages: int) -> list[str]:
    html = fetch_url(url)
    soup = BeautifulSoup(html, "lxml")
    base = urlparse(url)

    links: list[str] = []
    for tag in soup.find_all("a", href=True):
        href = urljoin(url, tag["href"])
        href, _ = urldefrag(href)
        parsed = urlparse(href)
        if parsed.scheme not in {"http", "https"}:
            continue
        if parsed.netloc != base.netloc:
            continue
        if allowlist and allowlist not in parsed.path:
            continue
        links.append(href)
        if len(links) >= max_pages:
            break

    return links


def fetch_url(url: str) -> str:
    resp = requests.get(
        url,
        timeout=20,
        headers={"User-Agent": "ai-tutor-rag/0.1"},
    )
    resp.raise_for_status()
    return resp.text


def extract_text_from_html(html: str) -> str:
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()

    main = soup.find("main")
    text = main.get_text(separator="\n") if main else soup.get_text(separator="\n")
    cleaned = normalize_whitespace(text)
    return cleaned


def extract_pdf_pages(path: Path) -> list[tuple[int, str]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[tuple[int, str]] = []
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        cleaned = normalize_whitespace(text)
        if cleaned:
            pages.append((idx, cleaned))
    return pages


def extract_text_from_pdf(path: Path) -> str:
    pages = extract_pdf_pages(path)
    return "\n".join(text for _, text in pages)


def extract_text_from_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return normalize_whitespace("\n".join(parts))


def extract_title_from_html(html: str) -> str | None:
    soup = BeautifulSoup(html, "lxml")
    if soup.title and soup.title.get_text(strip=True):
        return normalize_whitespace(soup.title.get_text())
    h1 = soup.find("h1")
    if h1 and h1.get_text(strip=True):
        return normalize_whitespace(h1.get_text())
    return None


def load_document(item: SourceItem) -> Document | None:
    if item.url:
        html = fetch_url(item.url)
        text = extract_text_from_html(html)
        if not text:
            return None
        title = extract_title_from_html(html) or item.title
        return Document(
            source_id=_source_id(item.url),
            title=title,
            source_url=item.url,
            text=text,
        )

    if item.path:
        path = Path(item.path)
        if not path.exists():
            return None
        suffix = path.suffix.lower()
        if suffix in {".html", ".htm"}:
            raw = path.read_text(encoding="utf-8", errors="ignore")
            text = extract_text_from_html(raw)
        elif suffix == ".pdf":
            text = extract_text_from_pdf(path)
        elif suffix == ".docx":
            text = extract_text_from_docx(path)
        else:
            raw = path.read_text(encoding="utf-8", errors="ignore")
            text = normalize_whitespace(raw)
        if not text:
            return None
        source_url = item.path
        return Document(
            source_id=_source_id(item.path),
            title=item.title,
            source_url=source_url,
            text=text,
        )

    return None


def build_chunks(
    documents: Iterable[Document],
    chunk_size: int = 1200,
    overlap: int = 200,
) -> list[Chunk]:
    chunks: list[Chunk] = []
    for doc in documents:
        for idx, chunk in enumerate(chunk_text(doc.text, chunk_size, overlap)):
            chunk_id = _chunk_id(doc.source_id, idx)
            chunks.append(
                Chunk(
                    chunk_id=chunk_id,
                    source_id=doc.source_id,
                    title=doc.title,
                    source_url=doc.source_url,
                    text=chunk,
                )
            )
    return chunks


def _source_id(value: str) -> str:
    return hashlib.sha1(value.encode("utf-8")).hexdigest()


def _chunk_id(source_id: str, idx: int) -> str:
    raw = f"{source_id}:{idx}"
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()
